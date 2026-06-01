# Copyright (c) 2024-2026 T1 Agentics LLC
# SPDX-License-Identifier: Apache-2.0

"""
Knowledge Base Service

Manages the Company Best Practices Database (SOP knowledge base).
Provides CRUD operations, search, and AI-powered queries for:
- Standard Operating Procedures (SOPs)
- Incident Response Playbooks
- Escalation Policies
- Compliance Requirements
- Agent Permission Boundaries
- Response Action Approvals

Features:
- Document parsing (PDF, DOCX, TXT, MD)
- Semantic search via embeddings
- Full-text search with PostgreSQL
- AI-powered document summarization
- Agent context injection

This knowledge base is queryable by AI agents during investigations.
"""

import json
import logging
import uuid
import os
import re
import hashlib
import io
from datetime import datetime
from typing import Dict, Any, Optional, List, Tuple
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """Result of document parsing"""
    content: str
    title: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)
    page_count: int = 0
    word_count: int = 0
    sections: List[Dict[str, str]] = field(default_factory=list)
    error: Optional[str] = None


@dataclass
class SemanticSearchResult:
    """Result from semantic search"""
    kb_id: str
    title: str
    content_snippet: str
    similarity_score: float
    content_type: str
    category: Optional[str] = None


# Content types for the knowledge base
CONTENT_TYPES = [
    'sop',              # Standard Operating Procedure
    'playbook',         # Incident Response Playbook
    'escalation',       # Escalation Policy
    'compliance',       # Compliance Requirement
    'permission',       # Agent Permission Boundary
    'approval_rule',    # Response Action Approval Rule
    'handling_rule',    # Alert Handling Rule
    'runbook',          # Technical Runbook
    'policy',           # General Policy
    'procedure',        # General Procedure
]

# Categories for organizing knowledge base entries
CATEGORIES = [
    'incident_response',
    'threat_detection',
    'malware_analysis',
    'phishing',
    'data_loss',
    'insider_threat',
    'network_security',
    'endpoint_security',
    'cloud_security',
    'identity_access',
    'compliance',
    'escalation',
    'communication',
    'documentation',
    'integrations',
    'general',
]


class KnowledgeBaseService:
    """
    Service for managing the Company Best Practices knowledge base.

    Features:
    - CRUD operations for knowledge base entries
    - Full-text search with filtering
    - Version control for changes
    - AI-powered document processing
    - Query interface for AI agents
    """

    def __init__(self):
        self.enabled = True

    async def create_entry(
        self,
        title: str,
        content: str,
        content_type: str = 'sop',
        category: Optional[str] = None,
        subcategory: Optional[str] = None,
        tags: List[str] = None,
        severity_filter: List[str] = None,
        incident_types: List[str] = None,
        ioc_types: List[str] = None,
        mitre_techniques: List[str] = None,
        compliance_frameworks: List[str] = None,
        priority: int = 100,
        created_by: str = 'system',
        source_document_name: Optional[str] = None,
        source_document_type: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Create a new knowledge base entry.

        Args:
            title: Entry title
            content: Full content/text of the entry
            content_type: Type of content (sop, playbook, escalation, etc.)
            category: Primary category
            subcategory: Optional subcategory
            tags: List of tags for searchability
            severity_filter: Which severities this applies to (low, medium, high, critical)
            incident_types: Types of incidents this applies to
            ioc_types: IOC types this is relevant for (ip, domain, hash, etc.)
            mitre_techniques: MITRE ATT&CK technique IDs
            compliance_frameworks: Compliance frameworks (NIST, SOC2, ISO27001, etc.)
            priority: Priority for ordering (lower = higher priority)
            created_by: Username of creator
            source_document_name: Original document filename if uploaded
            source_document_type: Original document type if uploaded

        Returns:
            Created entry data
        """
        try:
            from services.postgres_db import postgres_db

            if not postgres_db.connected:
                return {"error": "Database not connected"}

            kb_id = f"KB-{uuid.uuid4().hex[:8].upper()}"

            # Generate embedding for semantic search
            embedding = None
            try:
                # Combine title and content for embedding (truncate to avoid token limits)
                embedding_text = f"{title}\n\n{content}"[:8000]
                embedding = await self._generate_embedding(embedding_text)
                if embedding:
                    logger.debug(f"Generated embedding for {kb_id} ({len(embedding)} dimensions)")
            except Exception as e:
                logger.warning(f"Failed to generate embedding for {kb_id}: {e}")

            async with postgres_db.tenant_acquire() as conn:
                row = await conn.fetchrow('''
                    INSERT INTO knowledge_base (
                        kb_id, title, content, content_type,
                        category, subcategory, tags, severity_filter,
                        incident_types, ioc_types, mitre_techniques,
                        compliance_frameworks, priority, created_by,
                        source_document_name, source_document_type, embedding
                    ) VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, $10, $11, $12, $13, $14, $15, $16, $17)
                    RETURNING *
                ''',
                    kb_id,
                    title,
                    content,
                    content_type,
                    category,
                    subcategory,
                    tags or [],
                    severity_filter or [],
                    incident_types or [],
                    ioc_types or [],
                    mitre_techniques or [],
                    compliance_frameworks or [],
                    priority,
                    created_by,
                    source_document_name,
                    source_document_type,
                    embedding
                )

                logger.info(f"Created knowledge base entry: {kb_id} (embedding: {'yes' if embedding else 'no'})")
                return self._row_to_dict(row)

        except Exception as e:
            logger.error(f"Failed to create knowledge base entry: {e}")
            return {"error": str(e)}

    async def get_entry(self, kb_id: str) -> Optional[Dict[str, Any]]:
        """Get a single knowledge base entry by ID."""
        try:
            from services.postgres_db import postgres_db

            if not postgres_db.connected:
                return None

            async with postgres_db.tenant_acquire() as conn:
                row = await conn.fetchrow(
                    'SELECT * FROM knowledge_base WHERE kb_id = $1',
                    kb_id
                )
                return self._row_to_dict(row) if row else None

        except Exception as e:
            logger.error(f"Failed to get knowledge base entry {kb_id}: {e}")
            return None

    async def list_entries(
        self,
        content_type: Optional[str] = None,
        category: Optional[str] = None,
        subcategory: Optional[str] = None,
        tags: Optional[List[str]] = None,
        severity: Optional[str] = None,
        incident_type: Optional[str] = None,
        ioc_type: Optional[str] = None,
        is_active: Optional[bool] = True,
        search_query: Optional[str] = None,
        source: Optional[str] = None,
        limit: int = 100,
        offset: int = 0
    ) -> Tuple[List[Dict[str, Any]], int]:
        """
        List knowledge base entries with filtering.

        Returns a tuple of (entries, total_count).
        Search ranking: title matches score highest, then tags, then content.
        """
        try:
            from services.postgres_db import postgres_db

            if not postgres_db.connected:
                return ([], 0)

            async with postgres_db.tenant_acquire() as conn:
                query_parts = ['SELECT * FROM knowledge_base WHERE 1=1']
                params = []
                param_count = 1

                if content_type:
                    query_parts.append(f'AND content_type = ${param_count}')
                    params.append(content_type)
                    param_count += 1

                if category:
                    query_parts.append(f'AND category = ${param_count}')
                    params.append(category)
                    param_count += 1

                if subcategory:
                    query_parts.append(f'AND subcategory = ${param_count}')
                    params.append(subcategory)
                    param_count += 1

                if tags:
                    query_parts.append(f'AND tags && ${param_count}')
                    params.append(tags)
                    param_count += 1

                if severity:
                    query_parts.append(f'AND ${param_count} = ANY(severity_filter)')
                    params.append(severity)
                    param_count += 1

                if incident_type:
                    query_parts.append(f'AND ${param_count} = ANY(incident_types)')
                    params.append(incident_type)
                    param_count += 1

                if ioc_type:
                    query_parts.append(f'AND ${param_count} = ANY(ioc_types)')
                    params.append(ioc_type)
                    param_count += 1

                if is_active is not None:
                    query_parts.append(f'AND is_active = ${param_count}')
                    params.append(is_active)
                    param_count += 1

                if source:
                    query_parts.append(f'AND source = ${param_count}')
                    params.append(source)
                    param_count += 1

                if search_query:
                    # Match on title, tags (as text), or content
                    like_param = param_count
                    tsq_param = param_count + 1
                    query_parts.append(f'''AND (
                        title ILIKE ${like_param}
                        OR array_to_string(tags, ' ') ILIKE ${like_param}
                        OR to_tsvector('english', title || ' ' || COALESCE(subcategory, '') || ' ' || content) @@ plainto_tsquery('english', ${tsq_param})
                        OR content ILIKE ${like_param}
                    )''')
                    params.append(f'%{search_query}%')
                    params.append(search_query)
                    param_count += 2

                    # Rank: title exact > title partial > tags > content
                    query_parts[0] = f'''SELECT *,
                        CASE
                            WHEN title ILIKE ${like_param} THEN 0
                            WHEN array_to_string(tags, ' ') ILIKE ${like_param} THEN 1
                            ELSE 2
                        END AS _search_rank
                        FROM knowledge_base WHERE 1=1'''
                    query_parts.append(f'ORDER BY _search_rank ASC, priority ASC, created_at DESC')
                else:
                    query_parts.append(f'ORDER BY priority ASC, created_at DESC')

                # Build count query from the same WHERE conditions (before ORDER/LIMIT)
                count_parts = [p for p in query_parts]
                count_parts[0] = 'SELECT COUNT(*) FROM knowledge_base WHERE 1=1'
                # Remove ORDER BY from count query
                count_parts = [p for p in count_parts if not p.startswith('ORDER BY')]
                count_query = ' '.join(count_parts)
                # params at this point has only the filter params, no limit/offset yet
                total_count = await conn.fetchval(count_query, *params)

                query_parts.append(f'LIMIT ${param_count} OFFSET ${param_count + 1}')
                params.extend([limit, offset])

                query = ' '.join(query_parts)
                rows = await conn.fetch(query, *params)

                results = [self._row_to_dict(row) for row in rows]
                # Attach total_count as attribute for the route to use
                return (results, total_count or 0)

        except Exception as e:
            logger.error(f"Failed to list knowledge base entries: {e}")
            return ([], 0)

    async def update_entry(
        self,
        kb_id: str,
        updates: Dict[str, Any],
        updated_by: str = 'system',
        change_reason: Optional[str] = None
    ) -> Optional[Dict[str, Any]]:
        """
        Update a knowledge base entry.
        Creates a version history record before updating.

        Args:
            kb_id: Entry ID
            updates: Dictionary of fields to update
            updated_by: Username of updater
            change_reason: Reason for the change

        Returns:
            Updated entry or None
        """
        try:
            from services.postgres_db import postgres_db

            if not postgres_db.connected:
                return None

            async with postgres_db.tenant_acquire() as conn:
                # Get current entry for versioning
                current = await conn.fetchrow(
                    'SELECT * FROM knowledge_base WHERE kb_id = $1',
                    kb_id
                )

                if not current:
                    return None

                # Save current version to history
                await conn.execute('''
                    INSERT INTO knowledge_base_versions (
                        kb_id, version, title, content, ai_summary,
                        ai_extracted_rules, changed_by, change_reason
                    ) VALUES ($1, $2, $3, $4, $5, $6, $7, $8)
                ''',
                    kb_id,
                    current['version'],
                    current['title'],
                    current['content'],
                    current['ai_summary'],
                    json.dumps(current['ai_extracted_rules'] or []),
                    updated_by,
                    change_reason
                )

                # Build update query
                allowed_fields = [
                    'title', 'content', 'content_type', 'category', 'subcategory',
                    'tags', 'severity_filter', 'incident_types', 'ioc_types',
                    'mitre_techniques', 'compliance_frameworks', 'priority',
                    'is_active', 'ai_processed', 'ai_summary', 'ai_extracted_rules'
                ]

                set_parts = ['version = version + 1', 'updated_by = $1', 'updated_at = CURRENT_TIMESTAMP']
                params = [updated_by]
                param_count = 2

                for field, value in updates.items():
                    if field in allowed_fields:
                        if field == 'ai_extracted_rules' and isinstance(value, (list, dict)):
                            value = json.dumps(value)
                        set_parts.append(f'{field} = ${param_count}')
                        params.append(value)
                        param_count += 1

                params.append(kb_id)

                query = f'''
                    UPDATE knowledge_base
                    SET {', '.join(set_parts)}
                    WHERE kb_id = ${param_count}
                    RETURNING *
                '''

                row = await conn.execute(query, *params)

                # Fetch updated row
                updated = await conn.fetchrow(
                    'SELECT * FROM knowledge_base WHERE kb_id = $1',
                    kb_id
                )

                logger.info(f"Updated knowledge base entry: {kb_id}")
                return self._row_to_dict(updated)

        except Exception as e:
            logger.error(f"Failed to update knowledge base entry {kb_id}: {e}")
            return None

    async def delete_entry(self, kb_id: str, deleted_by: str = 'system') -> dict:
        """
        Delete a knowledge base entry (soft delete by setting is_active=False).
        Builtin articles (source='builtin') cannot be deleted.

        Returns:
            dict with 'deleted' bool and optional 'error' string
        """
        try:
            from services.postgres_db import postgres_db

            if not postgres_db.connected:
                return {"deleted": False, "error": "Database not connected"}

            async with postgres_db.tenant_acquire() as conn:
                row = await conn.fetchrow(
                    "SELECT source FROM knowledge_base WHERE kb_id = $1", kb_id
                )
                if not row:
                    return {"deleted": False, "error": "not_found"}

                if row.get("source") == "builtin":
                    return {"deleted": False, "error": "builtin_protected"}

                result = await conn.execute('''
                    UPDATE knowledge_base
                    SET is_active = FALSE, updated_at = CURRENT_TIMESTAMP
                    WHERE kb_id = $1
                ''', kb_id)

                deleted = result == 'UPDATE 1'
                if deleted:
                    logger.info(f"Deleted knowledge base entry: {kb_id}")
                return {"deleted": deleted}

        except Exception as e:
            logger.error(f"Failed to delete knowledge base entry {kb_id}: {e}")
            return {"deleted": False, "error": str(e)}

    async def get_entry_versions(self, kb_id: str) -> List[Dict[str, Any]]:
        """Get version history for a knowledge base entry."""
        try:
            from services.postgres_db import postgres_db

            if not postgres_db.connected:
                return []

            async with postgres_db.tenant_acquire() as conn:
                rows = await conn.fetch('''
                    SELECT * FROM knowledge_base_versions
                    WHERE kb_id = $1
                    ORDER BY version DESC
                ''', kb_id)

                return [dict(row) for row in rows]

        except Exception as e:
            logger.error(f"Failed to get versions for {kb_id}: {e}")
            return []

    async def query_for_context(
        self,
        alert_data: Dict[str, Any] = None,
        investigation_data: Dict[str, Any] = None,
        severity: Optional[str] = None,
        incident_type: Optional[str] = None,
        ioc_types: Optional[List[str]] = None,
        mitre_techniques: Optional[List[str]] = None,
        keywords: Optional[List[str]] = None,
        limit: int = 10,
        alert_type: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """
        Query knowledge base for relevant context during AI investigation.

        This is the primary method used by AI agents to retrieve SOPs,
        playbooks, and handling rules relevant to the current investigation.

        Uses semantic search (vector similarity) when available, falls back to
        keyword matching if embeddings aren't available.

        Args:
            alert_data: Alert being investigated
            investigation_data: Investigation context
            severity: Severity level to filter by
            incident_type: Type of incident
            ioc_types: IOC types being analyzed
            mitre_techniques: MITRE techniques to match
            keywords: Additional search keywords
            limit: Max results to return
            alert_type: Type of alert (email, endpoint, network, generic) for filtering

        Returns:
            List of relevant knowledge base entries with content
        """
        try:
            from services.postgres_db import postgres_db

            if not postgres_db.connected:
                return []

            # Extract context from alert/investigation if provided
            if alert_data:
                severity = severity or alert_data.get('severity')
                if alert_data.get('raw_event'):
                    raw = alert_data['raw_event']
                    if isinstance(raw, str):
                        try:
                            raw = json.loads(raw)
                        except:
                            pass
                    # Try to extract incident type from alert
                    incident_type = incident_type or raw.get('category') or raw.get('incident_type')

            # ═══════════════════════════════════════════════════════════════════
            # STRATEGY 1: Semantic Search (Preferred)
            # Uses AI embeddings to find conceptually similar SOPs
            # ═══════════════════════════════════════════════════════════════════
            try:
                # Build semantic search query from all available context
                query_parts = []

                if alert_data:
                    if alert_data.get('title'):
                        query_parts.append(alert_data['title'])
                    if alert_data.get('description'):
                        query_parts.append(alert_data['description'])

                if keywords:
                    query_parts.extend(keywords)

                if incident_type:
                    query_parts.append(incident_type)

                # Combine into search query
                search_query = ' '.join(query_parts) if query_parts else None

                if search_query:
                    logger.debug(f"[KB_QUERY] Using semantic search: {search_query[:100]}")

                    # Use semantic search with lower threshold for triage (more permissive)
                    results = await self.semantic_search(
                        query=search_query,
                        limit=limit,
                        min_similarity=0.6,  # Lower threshold for triage context
                        content_types=['sop', 'playbook', 'procedure', 'runbook']
                    )

                    if results:
                        logger.info(f"[KB_QUERY] Semantic search found {len(results)} entries")
                        return results
                    else:
                        logger.debug("[KB_QUERY] Semantic search returned no results, falling back to keyword search")

            except Exception as semantic_err:
                logger.warning(f"[KB_QUERY] Semantic search failed ({semantic_err}), falling back to keyword search")

            # ═══════════════════════════════════════════════════════════════════
            # STRATEGY 2: Keyword Search (Fallback)
            # Traditional SQL ILIKE matching when semantic search unavailable
            # ═══════════════════════════════════════════════════════════════════
            async with postgres_db.tenant_acquire() as conn:
                query_parts = ['SELECT * FROM knowledge_base WHERE is_active = TRUE']
                params = []
                param_count = 1

                # Build OR conditions for matching
                or_conditions = []

                if severity:
                    or_conditions.append(f'${param_count} = ANY(severity_filter)')
                    params.append(severity)
                    param_count += 1

                if incident_type:
                    or_conditions.append(f'${param_count} = ANY(incident_types)')
                    params.append(incident_type)
                    param_count += 1

                if ioc_types:
                    for ioc_type in ioc_types:
                        or_conditions.append(f'${param_count} = ANY(ioc_types)')
                        params.append(ioc_type)
                        param_count += 1

                if mitre_techniques:
                    for technique in mitre_techniques:
                        or_conditions.append(f'${param_count} = ANY(mitre_techniques)')
                        params.append(technique)
                        param_count += 1

                if keywords:
                    for keyword in keywords:
                        or_conditions.append(f"(title ILIKE ${param_count} OR content ILIKE ${param_count})")
                        params.append(f'%{keyword}%')
                        param_count += 1

                # If we have any conditions, add them
                if or_conditions:
                    query_parts.append(f'AND ({" OR ".join(or_conditions)})')

                query_parts.append('ORDER BY priority ASC, created_at DESC')
                query_parts.append(f'LIMIT ${param_count}')
                params.append(limit)

                query = ' '.join(query_parts)
                rows = await conn.fetch(query, *params)

                logger.info(f"[KB_QUERY] Keyword search found {len(rows)} entries")

                results = []
                for row in rows:
                    entry = self._row_to_dict(row)

                    # Alert type filtering - exclude SOPs that don't match alert type
                    if alert_type:
                        entry_title = (entry.get('title', '') or '').lower()
                        entry_category = (entry.get('category', '') or '').lower()
                        entry_content = (entry.get('content', '') or '').lower()[:500]  # Check first 500 chars
                        entry_tags = [t.lower() for t in entry.get('tags', [])]

                        # Define exclusion patterns for each alert type
                        if alert_type == 'endpoint':
                            # Exclude email-specific SOPs for endpoint alerts
                            email_indicators = ['email', 'phishing', 'spam', 'sender', 'recipient',
                                               'mailbox', 'inbox', 'sop-email', 'phish', 'bec']
                            is_email_sop = any(
                                ind in entry_title or ind in entry_category or
                                ind in entry_content or ind in entry_tags
                                for ind in email_indicators
                            )
                            # Only exclude if it's clearly an email SOP
                            if is_email_sop and 'endpoint' not in entry_title and 'malware' not in entry_title:
                                logger.debug(f"[KB_FILTER] Excluding email SOP '{entry.get('kb_id')}' for endpoint alert")
                                continue

                        elif alert_type == 'email':
                            # Exclude endpoint-specific SOPs for email alerts
                            endpoint_indicators = ['endpoint', 'process', 'registry', 'file hash',
                                                   'executable', 'malware detection', 'edr', 'av']
                            is_endpoint_sop = any(
                                ind in entry_title or ind in entry_category
                                for ind in endpoint_indicators
                            )
                            if is_endpoint_sop and 'email' not in entry_title and 'attachment' not in entry_title:
                                logger.debug(f"[KB_FILTER] Excluding endpoint SOP '{entry.get('kb_id')}' for email alert")
                                continue

                        elif alert_type == 'network':
                            # Exclude email-specific SOPs for network alerts
                            email_indicators = ['email', 'phishing', 'spam', 'sender', 'recipient']
                            is_email_sop = any(ind in entry_title or ind in entry_category for ind in email_indicators)
                            if is_email_sop:
                                logger.debug(f"[KB_FILTER] Excluding email SOP '{entry.get('kb_id')}' for network alert")
                                continue

                    # For context queries, include the full content and AI summary
                    results.append({
                        'kb_id': entry['kb_id'],
                        'title': entry['title'],
                        'content_type': entry['content_type'],
                        'category': entry['category'],
                        'content': entry['content'],
                        'ai_summary': entry.get('ai_summary'),
                        'ai_extracted_rules': entry.get('ai_extracted_rules', []),
                        'priority': entry['priority'],
                        'tags': entry.get('tags', []),
                        'severity_filter': entry.get('severity_filter', []),
                        'incident_types': entry.get('incident_types', []),
                    })

                return results

        except Exception as e:
            logger.error(f"Failed to query knowledge base for context: {e}")
            return []

    async def approve_entry(
        self,
        kb_id: str,
        approved_by: str
    ) -> Optional[Dict[str, Any]]:
        """Mark a knowledge base entry as approved."""
        try:
            from services.postgres_db import postgres_db

            if not postgres_db.connected:
                return None

            async with postgres_db.tenant_acquire() as conn:
                await conn.execute('''
                    UPDATE knowledge_base
                    SET approved_by = $1, approved_at = CURRENT_TIMESTAMP, updated_at = CURRENT_TIMESTAMP
                    WHERE kb_id = $2
                ''', approved_by, kb_id)

                row = await conn.fetchrow(
                    'SELECT * FROM knowledge_base WHERE kb_id = $1',
                    kb_id
                )

                logger.info(f"Approved knowledge base entry: {kb_id} by {approved_by}")
                return self._row_to_dict(row) if row else None

        except Exception as e:
            logger.error(f"Failed to approve knowledge base entry {kb_id}: {e}")
            return None

    async def get_stats(self, source: Optional[str] = None) -> Dict[str, Any]:
        """Get statistics about the knowledge base. Optionally scoped to a source."""
        try:
            from services.postgres_db import postgres_db

            if not postgres_db.connected:
                return {"error": "Database not connected"}

            async with postgres_db.tenant_acquire() as conn:
                # Build source filter clause
                source_filter = ''
                source_params = []
                if source:
                    source_filter = ' AND source = $1'
                    source_params = [source]

                # Total counts
                total = await conn.fetchval(
                    f'SELECT COUNT(*) FROM knowledge_base WHERE is_active = TRUE{source_filter}',
                    *source_params
                )

                # By content type
                type_counts = await conn.fetch(f'''
                    SELECT content_type, COUNT(*) as count
                    FROM knowledge_base
                    WHERE is_active = TRUE{source_filter}
                    GROUP BY content_type
                    ORDER BY count DESC
                ''', *source_params)

                # By category
                category_counts = await conn.fetch(f'''
                    SELECT category, COUNT(*) as count
                    FROM knowledge_base
                    WHERE is_active = TRUE AND category IS NOT NULL{source_filter}
                    GROUP BY category
                    ORDER BY count DESC
                ''', *source_params)

                # By subcategory (grouped by category)
                subcategory_counts = await conn.fetch(f'''
                    SELECT category, subcategory, COUNT(*) as count
                    FROM knowledge_base
                    WHERE is_active = TRUE AND subcategory IS NOT NULL{source_filter}
                    GROUP BY category, subcategory
                    ORDER BY category, count DESC
                ''', *source_params)

                # Build subcategories map: {category: [{subcategory, count}, ...]}
                subcategories_by_category = {}
                for row in subcategory_counts:
                    cat = row['category']
                    if cat not in subcategories_by_category:
                        subcategories_by_category[cat] = []
                    subcategories_by_category[cat].append({
                        'subcategory': row['subcategory'],
                        'count': row['count']
                    })

                # Pending approval
                pending_approval = await conn.fetchval(
                    f'SELECT COUNT(*) FROM knowledge_base WHERE is_active = TRUE AND approved_at IS NULL{source_filter}',
                    *source_params
                )

                # AI processed
                ai_processed = await conn.fetchval(
                    f'SELECT COUNT(*) FROM knowledge_base WHERE is_active = TRUE AND ai_processed = TRUE{source_filter}',
                    *source_params
                )

                return {
                    'total_entries': total,
                    'by_content_type': {row['content_type']: row['count'] for row in type_counts},
                    'by_category': {row['category']: row['count'] for row in category_counts if row['category']},
                    'subcategories_by_category': subcategories_by_category,
                    'pending_approval': pending_approval,
                    'ai_processed': ai_processed,
                    'content_types': CONTENT_TYPES,
                    'categories': CATEGORIES,
                }

        except Exception as e:
            logger.error(f"Failed to get knowledge base stats: {e}")
            return {"error": str(e)}

    def _row_to_dict(self, row) -> Dict[str, Any]:
        """Convert database row to dictionary."""
        if not row:
            return None

        result = dict(row)

        # Remove internal search rank column if present
        result.pop('_search_rank', None)

        # Convert UUID to string
        if result.get('id'):
            result['id'] = str(result['id'])

        # Convert datetime to ISO string
        for field in ['created_at', 'updated_at', 'approved_at']:
            if result.get(field):
                result[field] = result[field].isoformat()

        # Parse JSONB fields
        if result.get('ai_extracted_rules'):
            if isinstance(result['ai_extracted_rules'], str):
                try:
                    result['ai_extracted_rules'] = json.loads(result['ai_extracted_rules'])
                except:
                    pass

        return result

    # =========================================================================
    # DOCUMENT PARSING
    # =========================================================================

    async def parse_document(
        self,
        file_content: bytes,
        filename: str,
        file_type: Optional[str] = None
    ) -> ParsedDocument:
        """
        Parse a document file and extract text content.

        Supports: PDF, DOCX, TXT, MD, HTML

        Args:
            file_content: Raw file bytes
            filename: Original filename
            file_type: Optional explicit file type

        Returns:
            ParsedDocument with extracted content
        """
        if not file_type:
            file_type = os.path.splitext(filename)[1].lower().lstrip('.')

        try:
            if file_type == 'pdf':
                return await self._parse_pdf(file_content, filename)
            elif file_type in ('docx', 'doc'):
                return await self._parse_docx(file_content, filename)
            elif file_type in ('txt', 'text'):
                return await self._parse_text(file_content, filename)
            elif file_type in ('md', 'markdown'):
                return await self._parse_markdown(file_content, filename)
            elif file_type in ('html', 'htm'):
                return await self._parse_html(file_content, filename)
            else:
                return ParsedDocument(
                    content='',
                    error=f"Unsupported file type: {file_type}"
                )
        except Exception as e:
            logger.error(f"Failed to parse document {filename}: {e}")
            return ParsedDocument(content='', error=str(e))

    async def _parse_pdf(self, file_content: bytes, filename: str) -> ParsedDocument:
        """Parse PDF document"""
        try:
            import pypdf
        except ImportError:
            # Fallback if pypdf not installed
            return ParsedDocument(
                content='',
                error="PDF parsing requires 'pypdf' package. Install with: pip install pypdf"
            )

        try:
            pdf_file = io.BytesIO(file_content)
            reader = pypdf.PdfReader(pdf_file)

            pages_text = []
            sections = []

            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ''
                pages_text.append(text)

                # Try to extract section headers
                lines = text.split('\n')
                for line in lines:
                    line = line.strip()
                    # Heuristic: short lines in all caps or starting with numbers might be headers
                    if len(line) > 3 and len(line) < 100:
                        if line.isupper() or re.match(r'^\d+\.?\s+', line):
                            sections.append({
                                'title': line,
                                'page': i + 1
                            })

            full_content = '\n\n'.join(pages_text)

            # Extract metadata
            metadata = {}
            if reader.metadata:
                if reader.metadata.title:
                    metadata['title'] = reader.metadata.title
                if reader.metadata.author:
                    metadata['author'] = reader.metadata.author
                if reader.metadata.creation_date:
                    metadata['created'] = str(reader.metadata.creation_date)

            return ParsedDocument(
                content=full_content,
                title=metadata.get('title') or os.path.splitext(filename)[0],
                metadata=metadata,
                page_count=len(reader.pages),
                word_count=len(full_content.split()),
                sections=sections[:20]  # Limit to first 20 sections
            )

        except Exception as e:
            return ParsedDocument(content='', error=f"PDF parsing error: {e}")

    async def _parse_docx(self, file_content: bytes, filename: str) -> ParsedDocument:
        """Parse DOCX document"""
        try:
            from docx import Document
        except ImportError:
            return ParsedDocument(
                content='',
                error="DOCX parsing requires 'python-docx' package. Install with: pip install python-docx"
            )

        try:
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)

            paragraphs = []
            sections = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

                    # Check if it's a heading
                    if para.style and 'Heading' in para.style.name:
                        sections.append({
                            'title': text,
                            'style': para.style.name
                        })

            full_content = '\n\n'.join(paragraphs)

            # Extract core properties
            metadata = {}
            try:
                if doc.core_properties.title:
                    metadata['title'] = doc.core_properties.title
                if doc.core_properties.author:
                    metadata['author'] = doc.core_properties.author
            except:
                pass

            return ParsedDocument(
                content=full_content,
                title=metadata.get('title') or os.path.splitext(filename)[0],
                metadata=metadata,
                page_count=0,  # DOCX doesn't have fixed pages
                word_count=len(full_content.split()),
                sections=sections[:20]
            )

        except Exception as e:
            return ParsedDocument(content='', error=f"DOCX parsing error: {e}")

    async def _parse_text(self, file_content: bytes, filename: str) -> ParsedDocument:
        """Parse plain text document"""
        try:
            # Try UTF-8 first, then fallback to latin-1
            try:
                content = file_content.decode('utf-8')
            except UnicodeDecodeError:
                content = file_content.decode('latin-1')

            # Extract sections from plain text (lines that look like headers)
            sections = []
            lines = content.split('\n')
            for i, line in enumerate(lines):
                line = line.strip()
                # Heuristic for headers: short lines followed by empty line or underlines
                if len(line) > 3 and len(line) < 80:
                    if i + 1 < len(lines):
                        next_line = lines[i + 1].strip()
                        if not next_line or set(next_line) <= {'=', '-', '_'}:
                            sections.append({'title': line})

            return ParsedDocument(
                content=content,
                title=os.path.splitext(filename)[0],
                word_count=len(content.split()),
                sections=sections[:20]
            )

        except Exception as e:
            return ParsedDocument(content='', error=f"Text parsing error: {e}")

    async def _parse_markdown(self, file_content: bytes, filename: str) -> ParsedDocument:
        """Parse Markdown document"""
        try:
            content = file_content.decode('utf-8')
        except UnicodeDecodeError:
            content = file_content.decode('latin-1')

        # Extract headers from markdown
        sections = []
        for match in re.finditer(r'^(#+)\s+(.+)$', content, re.MULTILINE):
            level = len(match.group(1))
            title = match.group(2).strip()
            sections.append({
                'title': title,
                'level': level
            })

        # Extract title from first H1
        title = os.path.splitext(filename)[0]
        h1_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        if h1_match:
            title = h1_match.group(1).strip()

        return ParsedDocument(
            content=content,
            title=title,
            word_count=len(content.split()),
            sections=sections[:20]
        )

    async def _parse_html(self, file_content: bytes, filename: str) -> ParsedDocument:
        """Parse HTML document and extract text"""
        try:
            content = file_content.decode('utf-8')
        except UnicodeDecodeError:
            content = file_content.decode('latin-1')

        # Simple HTML tag removal (for basic cases)
        # Remove script and style content
        content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL | re.IGNORECASE)
        content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL | re.IGNORECASE)

        # Extract headers before removing all tags
        sections = []
        for match in re.finditer(r'<h(\d)[^>]*>([^<]+)</h\d>', content, re.IGNORECASE):
            level = int(match.group(1))
            title = match.group(2).strip()
            sections.append({
                'title': title,
                'level': level
            })

        # Remove all HTML tags
        text = re.sub(r'<[^>]+>', ' ', content)
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text).strip()

        # Extract title from <title> tag
        title = os.path.splitext(filename)[0]
        title_match = re.search(r'<title>([^<]+)</title>', content, re.IGNORECASE)
        if title_match:
            title = title_match.group(1).strip()

        return ParsedDocument(
            content=text,
            title=title,
            word_count=len(text.split()),
            sections=sections[:20]
        )

    async def upload_and_parse_document(
        self,
        file_content: bytes,
        filename: str,
        created_by: str = 'system',
        content_type: str = 'sop',
        category: Optional[str] = None,
        tags: Optional[List[str]] = None,
        auto_summarize: bool = True
    ) -> Dict[str, Any]:
        """
        Upload a document, parse it, and create a knowledge base entry.

        Args:
            file_content: Raw file bytes
            filename: Original filename
            created_by: Username of uploader
            content_type: Type of content
            category: Category for the entry
            tags: Tags for searchability
            auto_summarize: Whether to auto-generate AI summary

        Returns:
            Created knowledge base entry or error
        """
        # Parse the document
        parsed = await self.parse_document(file_content, filename)

        if parsed.error:
            return {"error": parsed.error}

        if not parsed.content.strip():
            return {"error": "Document appears to be empty"}

        # Create hash of content for deduplication
        content_hash = hashlib.sha256(parsed.content.encode()).hexdigest()[:16]

        # Check for duplicate
        from services.postgres_db import postgres_db
        if postgres_db.connected:
            async with postgres_db.tenant_acquire() as conn:
                existing = await conn.fetchval('''
                    SELECT kb_id FROM knowledge_base
                    WHERE source_document_name = $1 AND is_active = TRUE
                ''', filename)

                if existing:
                    return {
                        "error": f"Document already exists as {existing}",
                        "existing_kb_id": existing
                    }

        # Create the entry
        entry = await self.create_entry(
            title=parsed.title or filename,
            content=parsed.content,
            content_type=content_type,
            category=category,
            tags=tags or [],
            created_by=created_by,
            source_document_name=filename,
            source_document_type=os.path.splitext(filename)[1].lower().lstrip('.')
        )

        if entry.get('error'):
            return entry

        # Add parsing metadata
        entry['parsing_info'] = {
            'page_count': parsed.page_count,
            'word_count': parsed.word_count,
            'sections': parsed.sections,
            'metadata': parsed.metadata
        }

        # Auto-summarize if requested
        if auto_summarize and entry.get('kb_id'):
            try:
                summary_result = await self.ai_summarize_entry(entry['kb_id'])
                if summary_result.get('success'):
                    entry['ai_summary'] = summary_result.get('summary')
            except Exception as e:
                logger.warning(f"Auto-summarize failed: {e}")

        return entry

    # =========================================================================
    # SEMANTIC SEARCH
    # =========================================================================

    async def semantic_search(
        self,
        query: str,
        limit: int = 10,
        content_types: Optional[List[str]] = None,
        categories: Optional[List[str]] = None,
        min_similarity: float = 0.5
    ) -> List[SemanticSearchResult]:
        """
        Perform semantic search using embeddings.

        This provides better search results than keyword matching by
        understanding the meaning of the query.

        Args:
            query: Search query
            limit: Max results
            content_types: Filter by content types
            categories: Filter by categories
            min_similarity: Minimum similarity threshold (0-1)

        Returns:
            List of SemanticSearchResult sorted by similarity
        """
        try:
            from services.postgres_db import postgres_db

            if not postgres_db.connected:
                return []

            # Generate embedding for query
            query_embedding = await self._generate_embedding(query)
            if not query_embedding:
                # Fallback to full-text search if embedding fails
                logger.warning("Embedding generation failed, falling back to full-text search")
                return await self._fallback_text_search(query, limit, content_types, categories)

            async with postgres_db.tenant_acquire() as conn:
                # Check if embeddings column exists
                has_embeddings = await conn.fetchval('''
                    SELECT EXISTS (
                        SELECT 1 FROM information_schema.columns
                        WHERE table_name = 'knowledge_base' AND column_name = 'embedding'
                    )
                ''')

                if not has_embeddings:
                    logger.warning("Embeddings column not found, falling back to full-text search")
                    return await self._fallback_text_search(query, limit, content_types, categories)

                # Build query with filters
                query_parts = ['''
                    SELECT kb_id, title, content, content_type, category,
                           1 - (embedding <=> $1::vector) as similarity
                    FROM knowledge_base
                    WHERE is_active = TRUE
                      AND embedding IS NOT NULL
                ''']
                params = [str(query_embedding)]
                param_count = 2

                if content_types:
                    query_parts.append(f'AND content_type = ANY(${param_count})')
                    params.append(content_types)
                    param_count += 1

                if categories:
                    query_parts.append(f'AND category = ANY(${param_count})')
                    params.append(categories)
                    param_count += 1

                query_parts.append(f'AND 1 - (embedding <=> $1::vector) >= ${param_count}')
                params.append(min_similarity)
                param_count += 1

                query_parts.append('ORDER BY similarity DESC')
                query_parts.append(f'LIMIT ${param_count}')
                params.append(limit)

                rows = await conn.fetch(' '.join(query_parts), *params)

                results = []
                for row in rows:
                    # Create content snippet
                    content = row['content'] or ''
                    snippet = content[:300] + '...' if len(content) > 300 else content

                    results.append(SemanticSearchResult(
                        kb_id=row['kb_id'],
                        title=row['title'],
                        content_snippet=snippet,
                        similarity_score=float(row['similarity']),
                        content_type=row['content_type'],
                        category=row['category']
                    ))

                return results

        except Exception as e:
            logger.error(f"Semantic search failed: {e}")
            return await self._fallback_text_search(query, limit, content_types, categories)

    async def _generate_embedding(self, text: str) -> Optional[List[float]]:
        """
        Generate embedding vector for text.

        Resolution order:
          1. Per-tenant BYO config (tenant_ai_config) — if 'disabled', skip
             embeddings entirely and let the FTS fallback serve search for
             this tenant.
          2. Platform OpenAI key (OPENAI_API_KEY env)
          3. Local sentence-transformers (MiniLM) fallback

        BYO embeddings must match the KB column width (VECTOR(1536)) — if
        the configured model produces a different dimension count, we log
        and return None so the row gets stored without embedding and the
        FTS path serves search cleanly.
        """
        import os

        # Resolve per-tenant config
        ctx = None
        try:
            from middleware.tenant_middleware import current_tenant_id
            from services import ai_provider_resolver
            _tid = current_tenant_id.get()
            if _tid:
                ctx = await ai_provider_resolver.resolve_embeddings(str(_tid))
        except Exception:
            ctx = None

        if ctx is not None and ctx.mode == "disabled":
            # Tenant explicitly turned embeddings off — FTS still works.
            return None

        try:
            import httpx

            if ctx is not None and ctx.mode == "byo":
                # BYO embeddings — OpenAI-compatible shape (both 'openai' and
                # 'self_hosted' speak the same /v1/embeddings API).
                base = (ctx.base_url or "https://api.openai.com").rstrip("/")
                url = f"{base}/v1/embeddings"
                headers = {"Content-Type": "application/json"}
                if ctx.api_key:
                    headers["Authorization"] = f"Bearer {ctx.api_key}"
                model_name = ctx.model or "text-embedding-3-small"

                async with httpx.AsyncClient() as client:
                    response = await client.post(
                        url,
                        headers=headers,
                        json={"input": text[:8000], "model": model_name},
                        timeout=30.0,
                    )
                if response.status_code != 200:
                    logger.warning(
                        f"BYO embedding call failed ({response.status_code}); "
                        f"skipping embedding for this row"
                    )
                    return None
                data = response.json()
                vec = data["data"][0]["embedding"]
                # Dimension safety — the KB column is VECTOR(1536). Other
                # widths can't be written without a column migration.
                expected_dim = ctx.dimensions or 1536
                if len(vec) != expected_dim:
                    logger.warning(
                        f"BYO embedding dimension mismatch: got {len(vec)}, "
                        f"expected {expected_dim}; skipping"
                    )
                    return None
                if len(vec) != 1536:
                    logger.warning(
                        f"BYO embedding dimension {len(vec)} != KB column width 1536; "
                        f"skipping (run migration to change column width)"
                    )
                    return None
                return vec

            # Platform path: OpenAI via env key
            openai_key = os.getenv('OPENAI_API_KEY')
            if openai_key:
                async with httpx.AsyncClient() as client:
                    response = await client.post(
                        'https://api.openai.com/v1/embeddings',
                        headers={
                            'Authorization': f'Bearer {openai_key}',
                            'Content-Type': 'application/json'
                        },
                        json={
                            'input': text[:8000],
                            'model': 'text-embedding-3-small'
                        },
                        timeout=30.0
                    )

                    if response.status_code == 200:
                        data = response.json()
                        return data['data'][0]['embedding']

            # Fallback: Try to use sentence-transformers locally
            try:
                from sentence_transformers import SentenceTransformer
                model = SentenceTransformer('all-MiniLM-L6-v2')
                embedding = model.encode(text[:8000]).tolist()
                # 384-dim model won't fit the 1536 column; log and skip
                # rather than write garbage.
                if len(embedding) != 1536:
                    logger.warning(
                        f"Local embedding fallback returns {len(embedding)} dims "
                        f"(KB column needs 1536); skipping embedding"
                    )
                    return None
                return embedding
            except ImportError:
                pass

            return None

        except Exception as e:
            logger.error(f"Failed to generate embedding: {e}")
            return None

    async def _fallback_text_search(
        self,
        query: str,
        limit: int,
        content_types: Optional[List[str]],
        categories: Optional[List[str]]
    ) -> List[SemanticSearchResult]:
        """
        Fallback to full-text search when embeddings unavailable.
        """
        try:
            from services.postgres_db import postgres_db

            if not postgres_db.connected:
                return []

            async with postgres_db.tenant_acquire() as conn:
                query_parts = ['''
                    SELECT kb_id, title, content, content_type, category,
                           ts_rank(to_tsvector('english', title || ' ' || content),
                                   plainto_tsquery('english', $1)) as rank
                    FROM knowledge_base
                    WHERE is_active = TRUE
                      AND to_tsvector('english', title || ' ' || content) @@ plainto_tsquery('english', $1)
                ''']
                params = [query]
                param_count = 2

                if content_types:
                    query_parts.append(f'AND content_type = ANY(${param_count})')
                    params.append(content_types)
                    param_count += 1

                if categories:
                    query_parts.append(f'AND category = ANY(${param_count})')
                    params.append(categories)
                    param_count += 1

                query_parts.append('ORDER BY rank DESC')
                query_parts.append(f'LIMIT ${param_count}')
                params.append(limit)

                rows = await conn.fetch(' '.join(query_parts), *params)

                results = []
                for row in rows:
                    content = row['content'] or ''
                    snippet = content[:300] + '...' if len(content) > 300 else content

                    # Convert rank to similarity-like score (0-1)
                    similarity = min(1.0, float(row['rank']) / 10)

                    results.append(SemanticSearchResult(
                        kb_id=row['kb_id'],
                        title=row['title'],
                        content_snippet=snippet,
                        similarity_score=similarity,
                        content_type=row['content_type'],
                        category=row['category']
                    ))

                return results

        except Exception as e:
            logger.error(f"Fallback text search failed: {e}")
            return []

    async def ai_summarize_entry(self, kb_id: str) -> Dict[str, Any]:
        """
        Use AI to generate a summary and extract rules from a KB entry.

        The summary is stored in ai_summary and extracted rules in ai_extracted_rules.
        """
        try:
            entry = await self.get_entry(kb_id)
            if not entry:
                return {"success": False, "error": "Entry not found"}

            content = entry.get('content', '')
            if len(content) < 50:
                return {"success": False, "error": "Content too short to summarize"}

            # Try to use Claude for summarization
            summary = None
            rules = []

            try:
                import os
                import httpx

                anthropic_key = os.getenv('ANTHROPIC_API_KEY')
                if anthropic_key:
                    async with httpx.AsyncClient() as client:
                        response = await client.post(
                            'https://api.anthropic.com/v1/messages',
                            headers={
                                'x-api-key': anthropic_key,
                                'anthropic-version': '2023-06-01',
                                'Content-Type': 'application/json'
                            },
                            json={
                                'model': 'claude-3-haiku-20240307',
                                'max_tokens': 1024,
                                'messages': [{
                                    'role': 'user',
                                    'content': f"""Analyze this security document and provide:
1. A 2-3 sentence summary
2. Key rules or procedures (as a JSON array of strings)

Document:
{content[:4000]}

Respond in this exact JSON format:
{{"summary": "...", "rules": ["rule1", "rule2", ...]}}"""
                                }]
                            },
                            timeout=60.0
                        )

                        if response.status_code == 200:
                            data = response.json()
                            text = data['content'][0]['text']
                            # Parse JSON from response
                            json_match = re.search(r'\{[^}]+\}', text, re.DOTALL)
                            if json_match:
                                parsed = json.loads(json_match.group())
                                summary = parsed.get('summary')
                                rules = parsed.get('rules', [])

            except Exception as e:
                logger.warning(f"AI summarization API failed: {e}")

            # Fallback: Simple extractive summary
            if not summary:
                sentences = content.split('.')[:3]
                summary = '. '.join(s.strip() for s in sentences if s.strip()) + '.'

            if not rules:
                # Extract lines that look like rules
                for line in content.split('\n'):
                    line = line.strip()
                    if re.match(r'^[-•*]\s', line) or re.match(r'^\d+\.\s', line):
                        rules.append(re.sub(r'^[-•*\d.]+\s*', '', line))
                        if len(rules) >= 10:
                            break

            # Update the entry
            await self.update_entry(
                kb_id,
                {
                    'ai_summary': summary,
                    'ai_extracted_rules': rules,
                    'ai_processed': True
                },
                updated_by='system'
            )

            return {
                "success": True,
                "kb_id": kb_id,
                "summary": summary,
                "rules": rules
            }

        except Exception as e:
            logger.error(f"AI summarization failed: {e}")
            return {"success": False, "error": str(e)}

    async def get_context_for_agent(
        self,
        alert_type: Optional[str] = None,
        ioc_types: Optional[List[str]] = None,
        severity: Optional[str] = None,
        keywords: Optional[List[str]] = None,
        max_entries: int = 5,
        max_content_length: int = 2000
    ) -> str:
        """
        Get formatted context from knowledge base for AI agent injection.

        Returns a formatted string suitable for including in agent prompts.

        Args:
            alert_type: Type of alert being investigated
            ioc_types: IOC types present in the investigation
            severity: Alert severity
            keywords: Additional keywords to search for
            max_entries: Maximum entries to include
            max_content_length: Max content per entry

        Returns:
            Formatted string with relevant SOPs and rules
        """
        results = await self.query_for_context(
            severity=severity,
            incident_type=alert_type,
            ioc_types=ioc_types,
            keywords=keywords,
            limit=max_entries
        )

        if not results:
            return ""

        context_parts = ["## Relevant SOPs and Guidelines\n"]

        for entry in results:
            content = entry.get('content', '')
            if len(content) > max_content_length:
                content = content[:max_content_length] + '...'

            context_parts.append(f"\n### {entry['title']}")
            context_parts.append(f"Type: {entry['content_type']}")
            if entry.get('ai_summary'):
                context_parts.append(f"Summary: {entry['ai_summary']}")
            if entry.get('ai_extracted_rules'):
                rules = entry['ai_extracted_rules'][:5]  # Limit rules
                context_parts.append("Key Rules:")
                for rule in rules:
                    context_parts.append(f"  - {rule}")
            context_parts.append(f"\nContent:\n{content}")
            context_parts.append("---")

        return '\n'.join(context_parts)


    # =========================================================================
    # RIGGS-SPECIFIC METHODS
    # =========================================================================

    async def create_riggs_draft(
        self,
        title: str,
        content: str,
        related_alerts: List[str] = None,
        suggested_tags: List[str] = None,
        suggested_mitre: List[str] = None,
        category: Optional[str] = None,
        content_type: str = 'procedure'
    ) -> Dict[str, Any]:
        """
        Create a draft KB article authored by Riggs.
        These require human approval before becoming active.

        Args:
            title: Article title
            content: Full content
            related_alerts: Alert IDs that informed this article
            suggested_tags: Riggs's suggested tags
            suggested_mitre: Suggested MITRE techniques
            category: Suggested category
            content_type: Type of content

        Returns:
            Created draft entry
        """
        try:
            from services.postgres_db import postgres_db

            if not postgres_db.connected:
                return {"error": "Database not connected"}

            kb_id = f"KB-{uuid.uuid4().hex[:8].upper()}"

            async with postgres_db.tenant_acquire() as conn:
                row = await conn.fetchrow('''
                    INSERT INTO knowledge_base (
                        kb_id, title, content, content_type, category,
                        tags, mitre_techniques, author_type, status,
                        related_alerts, created_by, is_active
                    ) VALUES ($1, $2, $3, $4, $5, $6, $7, 'riggs', 'draft', $8, 'riggs', FALSE)
                    RETURNING *
                ''',
                    kb_id,
                    title,
                    content,
                    content_type,
                    category,
                    suggested_tags or [],
                    suggested_mitre or [],
                    related_alerts or []
                )

                logger.info(f"Riggs created draft KB article: {kb_id}")
                return self._row_to_dict(row)

        except Exception as e:
            logger.error(f"Failed to create Riggs draft: {e}")
            return {"error": str(e)}

    async def get_riggs_drafts(self, limit: int = 50) -> List[Dict[str, Any]]:
        """Get all draft articles created by Riggs pending human review."""
        try:
            from services.postgres_db import postgres_db

            if not postgres_db.connected:
                return []

            async with postgres_db.tenant_acquire() as conn:
                rows = await conn.fetch('''
                    SELECT * FROM knowledge_base
                    WHERE author_type = 'riggs' AND status = 'draft'
                    ORDER BY created_at DESC
                    LIMIT $1
                ''', limit)

                return [self._row_to_dict(row) for row in rows]

        except Exception as e:
            logger.error(f"Failed to get Riggs drafts: {e}")
            return []

    async def approve_riggs_draft(
        self,
        kb_id: str,
        approved_by: str,
        edits: Optional[Dict[str, Any]] = None
    ) -> Optional[Dict[str, Any]]:
        """
        Approve a Riggs-authored draft, making it active.

        Args:
            kb_id: KB entry ID
            approved_by: Username of approver
            edits: Optional edits to apply before approval

        Returns:
            Updated entry or None
        """
        try:
            from services.postgres_db import postgres_db

            if not postgres_db.connected:
                return None

            async with postgres_db.tenant_acquire() as conn:
                # Apply edits if provided
                if edits:
                    allowed_edits = ['title', 'content', 'category', 'tags', 'mitre_techniques']
                    set_parts = []
                    params = []
                    param_count = 1

                    for field, value in edits.items():
                        if field in allowed_edits:
                            set_parts.append(f'{field} = ${param_count}')
                            params.append(value)
                            param_count += 1

                    if set_parts:
                        params.append(kb_id)
                        await conn.execute(f'''
                            UPDATE knowledge_base
                            SET {', '.join(set_parts)}
                            WHERE kb_id = ${param_count}
                        ''', *params)

                # Approve the draft
                await conn.execute('''
                    UPDATE knowledge_base
                    SET status = 'published',
                        is_active = TRUE,
                        approved_by = $1,
                        approved_at = CURRENT_TIMESTAMP,
                        updated_at = CURRENT_TIMESTAMP
                    WHERE kb_id = $2
                ''', approved_by, kb_id)

                row = await conn.fetchrow(
                    'SELECT * FROM knowledge_base WHERE kb_id = $1',
                    kb_id
                )

                logger.info(f"Approved Riggs draft: {kb_id} by {approved_by}")
                return self._row_to_dict(row) if row else None

        except Exception as e:
            logger.error(f"Failed to approve Riggs draft {kb_id}: {e}")
            return None

    async def reject_riggs_draft(
        self,
        kb_id: str,
        rejected_by: str,
        reason: Optional[str] = None
    ) -> bool:
        """
        Reject a Riggs-authored draft.

        Args:
            kb_id: KB entry ID
            rejected_by: Username of rejector
            reason: Optional rejection reason

        Returns:
            True if rejected
        """
        try:
            from services.postgres_db import postgres_db

            if not postgres_db.connected:
                return False

            async with postgres_db.tenant_acquire() as conn:
                await conn.execute('''
                    UPDATE knowledge_base
                    SET status = 'archived',
                        is_active = FALSE,
                        updated_at = CURRENT_TIMESTAMP
                    WHERE kb_id = $1 AND author_type = 'riggs' AND status = 'draft'
                ''', kb_id)

                logger.info(f"Rejected Riggs draft: {kb_id} by {rejected_by}. Reason: {reason}")
                return True

        except Exception as e:
            logger.error(f"Failed to reject Riggs draft {kb_id}: {e}")
            return False

    async def record_kb_usage(self, kb_id: str) -> None:
        """
        Record that Riggs referenced a KB article.
        Used for tracking which articles are most valuable.
        """
        try:
            from services.postgres_db import postgres_db

            if not postgres_db.connected:
                return

            async with postgres_db.tenant_acquire() as conn:
                await conn.execute('''
                    UPDATE knowledge_base
                    SET usage_count = COALESCE(usage_count, 0) + 1,
                        last_used_at = CURRENT_TIMESTAMP
                    WHERE kb_id = $1
                ''', kb_id)

        except Exception as e:
            logger.error(f"Failed to record KB usage for {kb_id}: {e}")

    async def query_for_riggs(
        self,
        alert_data: Dict[str, Any],
        keywords: Optional[List[str]] = None,
        limit: int = 5
    ) -> List[Dict[str, Any]]:
        """
        Query KB for relevant context for Riggs analysis.
        Also records usage for referenced articles.

        Args:
            alert_data: Alert being analyzed
            keywords: Additional search keywords
            limit: Max results

        Returns:
            List of relevant KB articles with content
        """
        # Extract context from alert
        severity = alert_data.get('severity')
        title = alert_data.get('title', '')
        raw_event = alert_data.get('raw_event', {})

        if isinstance(raw_event, str):
            try:
                raw_event = json.loads(raw_event)
            except:
                raw_event = {}

        # Build keyword list
        search_keywords = keywords or []

        # Add keywords from alert title
        if title:
            search_keywords.extend(title.lower().split()[:5])

        # Get category/incident type
        incident_type = raw_event.get('category') or raw_event.get('incident_type')

        # Get MITRE techniques if present
        mitre = raw_event.get('mitre_techniques', [])
        if isinstance(mitre, str):
            mitre = [mitre]

        # Query KB
        results = await self.query_for_context(
            alert_data=alert_data,
            severity=severity,
            incident_type=incident_type,
            mitre_techniques=mitre,
            keywords=search_keywords,
            limit=limit
        )

        # Record usage for found articles
        for result in results:
            await self.record_kb_usage(result.get('kb_id'))

        return results

    async def semantic_search(
        self,
        query: str,
        limit: int = 10,
        min_similarity: float = 0.7,
        content_types: Optional[List[str]] = None,
        categories: Optional[List[str]] = None
    ) -> List[Dict[str, Any]]:
        """
        Perform semantic search using vector similarity.

        Args:
            query: Search query text
            limit: Maximum number of results
            min_similarity: Minimum similarity score (0-1)
            content_types: Filter by content types
            categories: Filter by categories

        Returns:
            List of similar KB entries with similarity scores
        """
        try:
            from services.postgres_db import postgres_db

            if not postgres_db.connected:
                return []

            # Generate embedding for the query
            query_embedding = await self._generate_embedding(query)
            if not query_embedding:
                logger.warning("Failed to generate embedding for semantic search")
                return []

            async with postgres_db.tenant_acquire() as conn:
                # Build filter conditions
                filter_conditions = ["kb.is_active = TRUE", "kb.embedding IS NOT NULL"]
                params = [query_embedding, min_similarity, limit]
                param_idx = 4

                if content_types:
                    filter_conditions.append(f"kb.content_type = ANY(${param_idx})")
                    params.append(content_types)
                    param_idx += 1

                if categories:
                    filter_conditions.append(f"kb.category = ANY(${param_idx})")
                    params.append(categories)
                    param_idx += 1

                where_clause = " AND ".join(filter_conditions)

                # Use cosine similarity (1 - cosine distance)
                sql_query = f"""
                    SELECT
                        kb.kb_id,
                        kb.title,
                        kb.content,
                        kb.content_type,
                        kb.category,
                        kb.tags,
                        kb.mitre_techniques,
                        kb.severity_filter,
                        kb.incident_types,
                        kb.priority,
                        kb.created_at,
                        kb.approved_at,
                        (1 - (kb.embedding <=> $1::vector)) AS similarity
                    FROM knowledge_base kb
                    WHERE {where_clause}
                        AND (1 - (kb.embedding <=> $1::vector)) >= $2
                    ORDER BY kb.embedding <=> $1::vector
                    LIMIT $3
                """

                rows = await conn.fetch(sql_query, *params)
                results = []
                for row in rows:
                    result = dict(row)
                    # Truncate content for preview
                    if result.get('content'):
                        result['content_snippet'] = result['content'][:500] + ('...' if len(result['content']) > 500 else '')
                    results.append(result)

                logger.info(f"Semantic search for '{query[:50]}...' found {len(results)} results")
                return results

        except Exception as e:
            logger.error(f"Semantic search failed: {e}", exc_info=True)
            return []


# Singleton instance
knowledge_base_service = KnowledgeBaseService()


def get_knowledge_base_service() -> KnowledgeBaseService:
    """Get the singleton knowledge base service instance."""
    return knowledge_base_service
